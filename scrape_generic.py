"""
Generic website scraper — extract readable page content from any URL and export to DOCX.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PlaywrightTimeoutError, sync_playwright

OUTPUT_DIR = Path(__file__).resolve().parent / "output"

EXTRACT_PAGE_JS = """
() => {
  const clean = (s) => (s || "").replace(/[ \\t]+/g, " ").replace(/\\n{3,}/g, "\\n\\n").trim();

  // Remove noisy elements before reading text
  const clone = document.body.cloneNode(true);
  clone.querySelectorAll(
    'script, style, noscript, svg, iframe, nav, footer, header, aside, ' +
    '[role="navigation"], [role="banner"], [aria-hidden="true"], ' +
    '.cookie, .cookies, #cookie, .advert, .ads, .sidebar'
  ).forEach((el) => el.remove());

  const meta = (name) => {
    const el =
      document.querySelector(`meta[name="${name}"]`) ||
      document.querySelector(`meta[property="${name}"]`) ||
      document.querySelector(`meta[property="og:${name}"]`);
    return el ? clean(el.getAttribute("content")) : "";
  };

  const title = clean(document.title || "");
  const h1 = clean((document.querySelector("h1") || {}).innerText || "");

  const headings = Array.from(document.querySelectorAll("h1, h2, h3"))
    .map((el) => ({
      level: el.tagName.toLowerCase(),
      text: clean(el.innerText || el.textContent || ""),
    }))
    .filter((h) => h.text && h.text.length < 300)
    .slice(0, 80);

  const mainEl =
    document.querySelector("article") ||
    document.querySelector("main") ||
    document.querySelector('[role="main"]') ||
    document.body;

  let mainText = clean(mainEl.innerText || mainEl.textContent || "");
  // Cap extremely long pages
  if (mainText.length > 120000) mainText = mainText.slice(0, 120000) + "\\n\\n[... truncated ...]";

  const links = [];
  const seenHref = new Set();
  for (const a of Array.from(document.querySelectorAll("a[href]"))) {
    const href = a.href || "";
    if (!href || href.startsWith("javascript:") || href === location.href + "#") continue;
    if (seenHref.has(href)) continue;
    seenHref.add(href);
    const text = clean(a.innerText || a.textContent || "");
    if (!text && !href) continue;
    links.push({ text: text.slice(0, 200), href });
    if (links.length >= 150) break;
  }

  const images = [];
  const seenSrc = new Set();
  for (const img of Array.from(document.querySelectorAll("img[src]"))) {
    const src = img.currentSrc || img.src || "";
    if (!src || seenSrc.has(src)) continue;
    if (src.startsWith("data:")) continue;
    seenSrc.add(src);
    images.push({
      alt: clean(img.alt || ""),
      src,
    });
    if (images.length >= 80) break;
  }

  return {
    scrape_type: "generic",
    title: title || h1 || location.hostname,
    h1,
    description: meta("description") || meta("og:description"),
    site_name: meta("og:site_name") || meta("application-name") || location.hostname,
    canonical: (document.querySelector('link[rel="canonical"]') || {}).href || location.href,
    url: location.href,
    headings,
    main_text: mainText,
    links,
    images,
    language: document.documentElement.lang || "",
  };
}
"""


def clean_text(value: str | None) -> str:
    if not value:
        return ""
    text = re.sub(r"[ \t]+", " ", value)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def dismiss_cookies(page: Page) -> None:
    for selector in [
        'button:has-text("Accept")',
        'button:has-text("Accept all")',
        'button:has-text("Accept All")',
        'button:has-text("Allow all")',
        'button:has-text("Allow All")',
        'button:has-text("Got it")',
        'button:has-text("I agree")',
        'button:has-text("Agree")',
        '[id*="accept" i]',
    ]:
        try:
            btn = page.locator(selector).first
            if btn.is_visible(timeout=600):
                btn.click(timeout=1200)
                page.wait_for_timeout(250)
                return
        except Exception:
            continue


def slug_from_url(url: str) -> str:
    parsed = urlparse(url)
    host = (parsed.hostname or "site").replace("www.", "")
    path = (parsed.path or "").strip("/").replace("/", "_")
    raw = f"{host}_{path}" if path else host
    slug = re.sub(r"[^a-zA-Z0-9._-]+", "_", raw).strip("_")
    return (slug or "page")[:80]


def add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value) if value else "—")


def build_docx(data: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(data.get("title") or "Web page scrape", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Exported {data.get('exported_at', '')} from {data.get('url', '')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Page info", level=1)
    add_kv(doc, "Title", data.get("title", ""))
    add_kv(doc, "H1", data.get("h1", ""))
    add_kv(doc, "Site", data.get("site_name", ""))
    add_kv(doc, "URL", data.get("url", ""))
    add_kv(doc, "Canonical", data.get("canonical", ""))
    add_kv(doc, "Language", data.get("language", ""))
    add_kv(doc, "Description", data.get("description", ""))

    headings = data.get("headings") or []
    if headings:
        doc.add_heading("Headings", level=1)
        for h in headings:
            level = h.get("level", "h2")
            prefix = {"h1": "#", "h2": "##", "h3": "###"}.get(level, "•")
            doc.add_paragraph(f"{prefix} {h.get('text', '')}")

    doc.add_heading("Main content", level=1)
    main_text = data.get("main_text") or "—"
    # Split into paragraphs for readability
    for block in re.split(r"\n\s*\n", main_text):
        block = clean_text(block)
        if block:
            doc.add_paragraph(block)

    links = data.get("links") or []
    if links:
        doc.add_heading(f"Links ({len(links)})", level=1)
        for link in links:
            text = link.get("text") or "(no text)"
            href = link.get("href") or ""
            doc.add_paragraph(f"{text}\n{href}")

    images = data.get("images") or []
    if images:
        doc.add_heading(f"Images ({len(images)})", level=1)
        for img in images:
            alt = img.get("alt") or "(no alt)"
            src = img.get("src") or ""
            doc.add_paragraph(f"{alt}\n{src}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def scrape(url: str, headless: bool = True, scroll_passes: int = 8) -> dict[str, Any]:
    if not re.match(r"^https?://", url.strip(), re.I):
        url = "https://" + url.strip()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        context = browser.new_context(
            viewport={"width": 1440, "height": 1600},
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            ),
            locale="en-US",
        )
        page = context.new_page()
        print(f"Opening {url}")
        page.goto(url.strip(), wait_until="domcontentloaded", timeout=90000)
        try:
            page.wait_for_load_state("networkidle", timeout=25000)
        except PlaywrightTimeoutError:
            pass
        page.wait_for_timeout(1500)
        dismiss_cookies(page)

        for _ in range(max(1, scroll_passes)):
            page.mouse.wheel(0, 1800)
            page.wait_for_timeout(300)
        page.evaluate("window.scrollTo(0, 0)")
        page.wait_for_timeout(400)

        # Click a few common "load more" / expand controls (best-effort)
        for sel in [
            'button:has-text("Show more")',
            'button:has-text("Load more")',
            'button:has-text("See more")',
            'a:has-text("Show more")',
            'a:has-text("Load more")',
        ]:
            try:
                loc = page.locator(sel)
                for i in range(min(loc.count(), 5)):
                    btn = loc.nth(i)
                    if btn.is_visible():
                        btn.click(timeout=1000)
                        page.wait_for_timeout(400)
            except Exception:
                continue

        print("Extracting page content...")
        data = page.evaluate(EXTRACT_PAGE_JS)
        data["exported_at"] = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        browser.close()
        return data
