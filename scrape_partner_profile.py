"""
Scrape a public Shopify Partner Directory profile and export all fields to DOCX.
Default: https://www.shopify.com/partners/directory/partner/forsuccess
"""

from __future__ import annotations

import argparse
import json
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlparse, urlunparse

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from playwright.sync_api import Page, TimeoutError as PlaywrightTimeoutError, sync_playwright

DEFAULT_URL = "https://www.shopify.com/partners/directory/partner/forsuccess"
OUTPUT_DIR = Path(__file__).resolve().parent / "output"


def clean_text(value: str | None) -> str:
    if not value:
        return ""
    text = re.sub(r"[ \t]+", " ", value)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def dismiss_cookies(page: Page) -> None:
    for selector in [
        'button:has-text("Accept")',
        'button:has-text("Accept all")',
        'button:has-text("Allow all")',
        'button:has-text("Got it")',
    ]:
        try:
            btn = page.locator(selector).first
            if btn.is_visible(timeout=800):
                btn.click(timeout=1500)
                page.wait_for_timeout(300)
                return
        except Exception:
            continue


def expand_truncated_reviews(page: Page) -> None:
    for _ in range(15):
        buttons = page.locator(
            '[data-section-name="reviews"] button:has-text("Show more"), '
            '[data-section-name="reviews"] a:has-text("Show more")'
        )
        count = buttons.count()
        if count == 0:
            break
        clicked = False
        for i in range(count):
            try:
                btn = buttons.nth(i)
                if btn.is_visible():
                    btn.click(timeout=1500)
                    clicked = True
                    page.wait_for_timeout(150)
            except Exception:
                continue
        if not clicked:
            break


EXTRACT_PROFILE_JS = """
() => {
  const text = (el) => (el ? (el.innerText || el.textContent || "").trim() : "");
  const clean = (s) => (s || "").replace(/[ \\t]+/g, " ").replace(/\\n{3,}/g, "\\n\\n").trim();
  const h1 = document.querySelector("h1");
  const main = document.querySelector("main") || document.body;
  const full = clean(text(main));

  let partnerSince = "";
  const sinceMatch = full.match(/Partner since\\s+([^\\n|]+)/i);
  if (sinceMatch) partnerSince = clean(sinceMatch[1]);

  let overallRating = "";
  let reviewCount = "";
  const ratingMatch = full.match(/(\\d(?:\\.\\d)?)\\s*\\((\\d[\\d,]*)\\)/);
  if (ratingMatch) {
    overallRating = ratingMatch[1];
    reviewCount = ratingMatch[2].replace(/,/g, "");
  }

  const labelValue = (label) => {
    const re = new RegExp(label + "\\\\s*\\\\n+([^\\\\n]+)", "i");
    const m = full.match(re);
    return m ? clean(m[1]) : "";
  };

  const emailMatch = full.match(/[A-Z0-9._%+-]+@[A-Z0-9.-]+\\.[A-Z]{2,}/i);
  const email = emailMatch ? emailMatch[0] : labelValue("Contact information");

  let businessDescription = "";
  const aboutIdx = full.indexOf("Business description");
  if (aboutIdx >= 0) {
    const after = full.slice(aboutIdx + "Business description".length);
    const stop = after.search(/\\nSpecialized services|\\nOther services|\\nIndustries|\\nFeatured work|\\nRating/i);
    businessDescription = clean(stop >= 0 ? after.slice(0, stop) : after.slice(0, 2000));
  }

  const specializedServices = [];
  const serviceNames = [
    "Store build or redesign",
    "Troubleshooting",
    "Theme customization",
    "Custom apps and integrations",
    "Store migration",
  ];
  for (const name of serviceNames) {
    const idx = full.indexOf(name);
    if (idx < 0) continue;
    const chunk = full.slice(idx, idx + 1200);
    const priceMatch = chunk.match(/Contact for pricing|\\$[\\d,]+(?:\\s*-\\s*\\$[\\d,]+)?/i);
    let description = "";
    const descMatch = chunk.match(/Description:\\s*([\\s\\S]*?)(?:Additional work and price information:|$)/i);
    if (descMatch) description = clean(descMatch[1]);
    let additional = "";
    const addMatch = chunk.match(/Additional work and price information:\\s*([\\s\\S]*?)(?=\\n(?:Store build or redesign|Troubleshooting|Theme customization|Custom apps and integrations|Store migration|Other services|Industries|Featured work|Rating)|$)/i);
    if (addMatch) additional = clean(addMatch[1]);
    specializedServices.push({
      name,
      pricing: priceMatch ? clean(priceMatch[0]) : "",
      description,
      additional_work: additional,
    });
  }

  let otherServices = [];
  const otherIdx = full.search(/Other services/i);
  if (otherIdx >= 0) {
    const after = full.slice(otherIdx + "Other services".length);
    const stop = after.search(/\\nMore services|\\nIndustries|\\nFeatured work|\\nRating/i);
    const block = clean(stop >= 0 ? after.slice(0, stop) : after.slice(0, 500));
    otherServices = block.split(/,|\\n/).map(clean).filter(Boolean).filter(s => !/^more services$/i.test(s));
  }

  let industries = [];
  const indIdx = full.search(/\\nIndustries\\n/i);
  if (indIdx >= 0) {
    const after = full.slice(indIdx).replace(/^\\n?Industries\\n/i, "");
    const stop = after.search(/\\nFeatured work|\\nRating/i);
    const block = clean(stop >= 0 ? after.slice(0, stop) : after.slice(0, 400));
    industries = block.split(/,/).map(clean).filter(Boolean);
  }

  const featuredWork = [];
  document.querySelectorAll("a, article, section").forEach((el) => {
    const t = clean(text(el));
    if (/View featured work/i.test(t) && t.length < 500) {
      const lines = t.split("\\n").map(clean).filter(Boolean).filter(l => !/View featured work/i.test(l));
      if (lines.length) {
        featuredWork.push({
          title: lines[0],
          description: lines.slice(1).join(" "),
          href: el.href || "",
        });
      }
    }
  });

  const ratingDistribution = {};
  const summaryIdx = full.search(/Overall rating summary/i);
  if (summaryIdx >= 0) {
    const summaryBlock = full.slice(summaryIdx, summaryIdx + 800);
    const counts = [...summaryBlock.matchAll(/\\((\\d+)\\)/g)].map(m => Number(m[1]));
    if (counts.length >= 5) {
      ratingDistribution["5_star"] = counts[0];
      ratingDistribution["4_star"] = counts[1];
      ratingDistribution["3_star"] = counts[2];
      ratingDistribution["2_star"] = counts[3];
      ratingDistribution["1_star"] = counts[4];
    }
  }

  let partnerType = "";
  if (/Service partner/i.test(full)) partnerType = "Service partner";
  else if (/Technology partner/i.test(full)) partnerType = "Technology partner";

  let priceRange = "";
  const priceIdx = full.search(/Price range for selected services/i);
  if (priceIdx >= 0) {
    const m = full.slice(priceIdx).match(/Price range for selected services\\s*\\n+([^\\n]+)/i);
    if (m) priceRange = clean(m[1]);
  }

  // Total review pages from pagination control
  let totalPages = 1;
  const pag = document.querySelector('[data-component-name="pagination"]');
  if (pag) {
    const t = pag.innerText || "";
    const m = t.match(/\\/\\s*(\\d+)/);
    if (m) totalPages = Number(m[1]);
  }

  return {
    name: clean(text(h1)),
    partner_type: partnerType,
    overall_rating: overallRating,
    review_count: reviewCount,
    partner_since: partnerSince,
    email,
    primary_location: labelValue("Primary location"),
    supported_locations: labelValue("Supported locations"),
    languages: labelValue("Languages"),
    price_range: priceRange,
    business_description: businessDescription,
    specialized_services: specializedServices,
    other_services: otherServices,
    industries,
    featured_work: featuredWork,
    rating_distribution: ratingDistribution,
    total_review_pages: totalPages,
    profile_url: location.href.split("?")[0],
  };
}
"""


EXTRACT_REVIEWS_JS = """
() => {
  const clean = (s) => (s || "").replace(/[ \\t]+/g, " ").replace(/\\n{3,}/g, "\\n\\n").trim();
  const section = document.querySelector('[data-section-name="reviews"]');
  if (!section) return [];

  const reviews = [];
  const seen = new Set();
  const blocks = Array.from(section.querySelectorAll("div")).filter((el) => {
    const t = el.innerText || "";
    if (!/Quality of work/i.test(t) || !/Communication/i.test(t)) return false;
    if (t.length < 40 || t.length > 5000) return false;
    const child = Array.from(el.querySelectorAll("div")).some((c) => {
      const ct = c.innerText || "";
      return (
        c !== el &&
        /Quality of work/i.test(ct) &&
        /Communication/i.test(ct) &&
        ct.length > 40 &&
        ct.length < t.length - 30
      );
    });
    return !child;
  });

  for (const el of blocks) {
    const t = clean(el.innerText || "");
    const lines = t.split("\\n").map(clean).filter(Boolean);
    if (lines.length < 4) continue;

    let company = lines[0];
    if (/^Reviews$/i.test(company) || /^Rating/i.test(company) || /^Overall/i.test(company)) continue;
    if (/^\\d+$/.test(company)) continue;

    let date = "";
    const dateLine = lines.find((l) =>
      /\\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\\.?\\s+\\d{1,2},\\s+\\d{4}\\b/i.test(l)
    );
    if (dateLine) {
      date = dateLine.match(
        /\\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\\.?\\s+\\d{1,2},\\s+\\d{4}\\b/i
      )[0];
    }

    let quality = "";
    let communication = "";
    const qIdx = lines.findIndex((l) => /^Quality of work$/i.test(l));
    if (qIdx >= 0 && lines[qIdx + 1] && /^[1-5]$/.test(lines[qIdx + 1])) quality = lines[qIdx + 1];
    const cIdx = lines.findIndex((l) => /^Communication$/i.test(l));
    if (cIdx >= 0 && lines[cIdx + 1] && /^[1-5]$/.test(lines[cIdx + 1])) communication = lines[cIdx + 1];

    let service = "";
    const svcLine = lines.find((l) => /Service reviewed:/i.test(l));
    if (svcLine) service = clean(svcLine.replace(/Service reviewed:\\s*/i, ""));

    const skip = new Set(
      [company, date, "Quality of work", quality, "Communication", communication, svcLine]
        .filter(Boolean)
        .map((x) => clean(x))
    );
    const bodyLines = lines.filter((l) => {
      if (skip.has(l)) return false;
      if (/^Quality of work$/i.test(l) || /^Communication$/i.test(l)) return false;
      if (/^Service reviewed:/i.test(l)) return false;
      if (/^Show more$/i.test(l) || /^Show less$/i.test(l)) return false;
      if (/^[1-5]$/.test(l)) return false;
      if (date && l.includes(date) && l.length < 40) return false;
      return true;
    });
    let body = clean(bodyLines.join("\\n"));
    body = body.replace(/\\s*Show more\\s*$/i, "").replace(/\\s*Show less\\s*$/i, "").trim();

    const key = `${company}|${date}|${body.slice(0, 80)}`;
    if (seen.has(key)) continue;
    seen.add(key);

    reviews.push({
      company_or_reviewer: company,
      date,
      quality_of_work: quality,
      communication,
      review_text: body,
      service_reviewed: service,
    });
  }
  return reviews;
}
"""


def current_review_page(page: Page) -> int:
    try:
        label = page.evaluate(
            """() => {
              const sel = document.querySelector('[data-section-name="reviews"] [role="option"][aria-selected="true"]');
              if (sel) {
                const m = (sel.getAttribute('aria-label') || '').match(/(\\d+)/);
                if (m) return Number(m[1]);
              }
              const span = document.querySelector('[data-section-name="reviews"] .pagination-select span');
              if (span && /^\\d+$/.test(span.textContent.trim())) return Number(span.textContent.trim());
              return 1;
            }"""
        )
        return int(label or 1)
    except Exception:
        return 1


def first_review_company(page: Page) -> str:
    return page.evaluate(
        """() => {
          const section = document.querySelector('[data-section-name="reviews"]');
          const text = section ? section.innerText : '';
          const lines = text.split('\\n').map(s => s.trim()).filter(Boolean);
          for (const l of lines) {
            if (l === 'Reviews') continue;
            if (/^\\d+$/.test(l)) continue;
            if (/Quality of work|Communication|Service reviewed|Show more|Show less/i.test(l)) continue;
            if (/^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)/i.test(l)) continue;
            return l;
          }
          return '';
        }"""
    )


def wait_for_reviews_loaded(page: Page, expected_page: int, previous_first: str | None) -> None:
    """Wait until reviews section shows content for expected_page."""
    deadline = time.time() + 30
    while time.time() < deadline:
        info = page.evaluate(
            """([expected, prev]) => {
              const section = document.querySelector('[data-section-name="reviews"]');
              const text = section ? section.innerText : '';
              const hasQuality = /Quality of work/i.test(text);
              const span = section && section.querySelector('.pagination-select span');
              const pageNum = span ? Number((span.textContent || '').trim()) : 0;
              const lines = text.split('\\n').map(s => s.trim()).filter(Boolean);
              let first = '';
              for (const l of lines) {
                if (l === 'Reviews') continue;
                if (/^\\d+$/.test(l)) continue;
                if (/Quality of work|Communication|Service reviewed|Show more|Show less/i.test(l)) continue;
                if (/^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)/i.test(l)) continue;
                first = l;
                break;
              }
              return {
                hasQuality,
                pageNum,
                first,
                ready:
                  hasQuality &&
                  pageNum === expected &&
                  (expected === 1 || !prev || first !== prev),
              };
            }""",
            [expected_page, previous_first or ""],
        )
        if info.get("ready"):
            return
        page.wait_for_timeout(400)


def click_next_review_page(page: Page) -> bool:
    nxt = page.locator('a[data-component-name="next-page"]')
    if not nxt.count():
        return False
    disabled = nxt.first.get_attribute("aria-disabled")
    if disabled == "true":
        return False
    before_page = current_review_page(page)
    before_first = first_review_company(page)
    nxt.first.click(force=True)
    wait_for_reviews_loaded(page, before_page + 1, before_first)
    # Extra settle time — review cards sometimes paint after the page marker updates
    page.wait_for_timeout(800)
    return current_review_page(page) == before_page + 1


def scrape_all_reviews(page: Page, total_pages: int) -> list[dict[str, Any]]:
    all_reviews: list[dict[str, Any]] = []
    seen: set[str] = set()

    def collect() -> int:
        expand_truncated_reviews(page)
        page.wait_for_timeout(250)
        batch = page.evaluate(EXTRACT_REVIEWS_JS)
        added = 0
        for r in batch:
            key = f"{r.get('company_or_reviewer')}|{r.get('date')}|{(r.get('review_text') or '')[:100]}"
            if key not in seen:
                seen.add(key)
                all_reviews.append(r)
                added += 1
        return added

    page.locator('[data-section-name="reviews"]').scroll_into_view_if_needed()
    page.wait_for_timeout(500)
    wait_for_reviews_loaded(page, 1, None)
    added = collect()
    print(f"  Page 1/{total_pages}: +{added} (total {len(all_reviews)})")

    for target in range(2, total_pages + 1):
        ok = click_next_review_page(page)
        if not ok:
            print(f"  Stopped: could not reach page {target} (on {current_review_page(page)})")
            break
        added = collect()
        print(f"  Page {target}/{total_pages}: +{added} (total {len(all_reviews)})")
        time.sleep(0.2)

    return all_reviews


def add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value) if value else "—")


def build_docx(data: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(data.get("name") or "Shopify Partner Profile", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Exported {data.get('exported_at', '')} from {data.get('profile_url', '')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Profile overview", level=1)
    add_kv(doc, "Partner type", data.get("partner_type", ""))
    add_kv(doc, "Overall rating", data.get("overall_rating", ""))
    add_kv(doc, "Review count", data.get("review_count", ""))
    add_kv(doc, "Partner since", data.get("partner_since", ""))
    add_kv(doc, "Email", data.get("email", ""))
    add_kv(doc, "Primary location", data.get("primary_location", ""))
    add_kv(doc, "Supported locations", data.get("supported_locations", ""))
    add_kv(doc, "Languages", data.get("languages", ""))
    add_kv(doc, "Price range", data.get("price_range", ""))

    doc.add_heading("About / Business description", level=1)
    doc.add_paragraph(data.get("business_description") or "—")

    doc.add_heading("Specialized services", level=1)
    for svc in data.get("specialized_services") or []:
        doc.add_heading(svc.get("name") or "Service", level=2)
        add_kv(doc, "Pricing", svc.get("pricing", ""))
        add_kv(doc, "Description", svc.get("description", ""))
        add_kv(doc, "Additional work and price information", svc.get("additional_work", ""))

    doc.add_heading("Other services", level=1)
    others = data.get("other_services") or []
    doc.add_paragraph(", ".join(others) if others else "—")

    doc.add_heading("Industries", level=1)
    industries = data.get("industries") or []
    doc.add_paragraph(", ".join(industries) if industries else "—")

    doc.add_heading("Featured work", level=1)
    for item in data.get("featured_work") or []:
        doc.add_heading(item.get("title") or "Project", level=2)
        if item.get("description"):
            doc.add_paragraph(item["description"])
        if item.get("href"):
            add_kv(doc, "Link", item["href"])

    doc.add_heading("Rating summary", level=1)
    dist = data.get("rating_distribution") or {}
    for star in ("5_star", "4_star", "3_star", "2_star", "1_star"):
        if star in dist:
            add_kv(doc, star.replace("_", " "), str(dist[star]))

    reviews = data.get("reviews") or []
    doc.add_heading(f"Reviews ({len(reviews)})", level=1)
    for i, rev in enumerate(reviews, 1):
        doc.add_heading(f"{i}. {rev.get('company_or_reviewer') or 'Review'}", level=2)
        add_kv(doc, "Date", rev.get("date", ""))
        add_kv(doc, "Quality of work", rev.get("quality_of_work", ""))
        add_kv(doc, "Communication", rev.get("communication", ""))
        add_kv(doc, "Service reviewed", rev.get("service_reviewed", ""))
        doc.add_paragraph(rev.get("review_text") or "—")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def scrape(url: str, headless: bool = True) -> dict[str, Any]:
    # Normalize to bare profile URL (no query)
    parsed = urlparse(url)
    url = urlunparse((parsed.scheme, parsed.netloc, parsed.path.rstrip("/"), "", "", ""))

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        context = browser.new_context(
            viewport={"width": 1440, "height": 1600},
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            ),
            locale="en-US",
        )
        page = context.new_page()
        print(f"Opening {url}")
        page.goto(url, wait_until="domcontentloaded", timeout=90000)
        try:
            page.wait_for_load_state("networkidle", timeout=30000)
        except PlaywrightTimeoutError:
            pass
        page.wait_for_timeout(2000)
        dismiss_cookies(page)

        for _ in range(5):
            page.mouse.wheel(0, 1600)
            page.wait_for_timeout(350)
        page.evaluate("window.scrollTo(0, 0)")
        page.wait_for_timeout(400)

        print("Extracting profile fields...")
        profile = page.evaluate(EXTRACT_PROFILE_JS)
        total_pages = int(profile.get("total_review_pages") or 1)
        print(f"Review pages detected: {total_pages}")

        try:
            page.locator('[data-section-name="reviews"]').scroll_into_view_if_needed(timeout=8000)
        except Exception:
            page.evaluate("window.scrollTo(0, document.body.scrollHeight * 0.75)")
        page.wait_for_timeout(600)

        print("Scraping reviews (all pages)...")
        reviews = scrape_all_reviews(page, total_pages)
        profile["reviews"] = reviews
        profile["reviews_scraped"] = len(reviews)
        profile["exported_at"] = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        browser.close()
        return profile


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export Shopify Partner Directory profile to DOCX. "
        "Prefer: python scrape.py --url <any-url> (auto-detects partner profiles)."
    )
    parser.add_argument("--url", default=DEFAULT_URL, help="Partner profile URL")
    parser.add_argument("--headed", action="store_true", help="Show browser window")
    parser.add_argument("--out", default="", help="Output DOCX path")
    args = parser.parse_args()

    data = scrape(args.url, headless=not args.headed)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    slug = args.url.rstrip("/").split("/")[-1].split("?")[0] or "partner"
    json_path = OUTPUT_DIR / f"{slug}_partner_profile.json"
    docx_path = Path(args.out) if args.out else OUTPUT_DIR / f"{slug}_partner_profile.docx"

    json_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    build_docx(data, docx_path)

    print("\nDone.")
    print(f"  Reviews scraped: {data.get('reviews_scraped', 0)} (listed count: {data.get('review_count')})")
    print(f"  JSON: {json_path}")
    print(f"  DOCX: {docx_path}")


if __name__ == "__main__":
    main()
